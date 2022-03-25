from __future__ import annotations
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches
from loguru import logger

if TYPE_CHECKING:
    from docx.document import Document as DocumentType


export_file_name = "export/demo.docx"

if __name__ == "__main__":

    document: DocumentType = Document()

    document.add_heading("Document Title", 0)

    p = document.add_paragraph("A plain paragraph having some ")
    p.add_run("bold").bold = True
    p.add_run(" and some ")
    p.add_run("italic.").italic = True

    document.add_heading("Heading, level 1", level=1)
    document.add_paragraph("Intense quote", style="Intense Quote")

    document.add_paragraph("first item in unordered list", style="List Bullet")
    document.add_paragraph("first item in ordered list", style="List Number")

    # add picture
    try:
        document.add_picture("monty-truth.png", width=Inches(1.25))
    except FileNotFoundError as err:
        logger.warning(err)

    # add table
    records = (
        (3, "101", "Spam"),
        (7, "422", "Eggs"),
        (4, "631", "Spam, spam, eggs, and spam"),
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Qty"
    hdr_cells[1].text = "Id"
    hdr_cells[2].text = "Desc"
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    try:
        document.save(export_file_name)
    except IOError as err:
        logger.error(f"failed write {export_file_name} see detail {err}")
    else:
        logger.info(f"succeed write file {export_file_name}")
